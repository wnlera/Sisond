import random
import re
import docx
from FormatChecker import ExtendedDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.document import Document

from xml.dom.minidom import parseString
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from typing import Union, List
from io import BytesIO

from FormatChecker.CheckResults import CheckResults
from FormatChecker import Mistakes, MistakeType, highlight_mistake
from FormatChecker.Utils.xml import *


def file_check_interface(file, selected_boxes, mock=True) -> CheckResults:
    file = BytesIO(file.read())
    if mock:
        return mock_return()
    return check_file(file, selected_boxes)

def mock_return():
    is_wrong = random.random() < 0.4
    res = [1 for x in range(5)]
    if is_wrong:
        for i in range(len(res)):
            if random.random() < 0.5:
                res[i] = 2
    results = CheckResults(res, BytesIO())
    return results


# ======================================================================
def check_file(file_like, selected_boxes):
    file = ExtendedDocument(file_like)

    verifiable_paras = []
    for para in file.docx.paragraphs[file.safe_table_of_content_index:]:
        verifiable_paras.append(para.text)  # todo: переделать, хранить копию документа не оч

    result = []
    # result = [file.table_of_content_index is not None, get_margin_is_ok(file.docx), get_font_is_ok(file),
    #           alignment_is_ok(file),
    #           line_spacing_is_ok(file), tables_is_ok(file.docx, verifiable_paras), pics_is_ok(file.docx, verifiable_paras)]
    if '0' in selected_boxes:
        result.append(file.table_of_content_index is not None)
    else:
        result.append(4)
    if '1' in selected_boxes:
        result.append(get_margin_is_ok(file.docx))
    else:
        result.append(4)
    if '2' in selected_boxes:
        result.append(get_font_is_ok(file))
    else:
        result.append(4)
    if '3' in selected_boxes:
        result.append(alignment_is_ok(file))
    else:
        result.append(4)
    if '4' in selected_boxes:
        result.append(line_spacing_is_ok(file))
    else:
        result.append(4)
    if '5' in selected_boxes:
        result.append(tables_is_ok(file.docx, verifiable_paras))
    else:
        result.append(4)
    if '6' in selected_boxes:
        result.append(pics_is_ok(file.docx, verifiable_paras))
    else:
        result.append(4)


    result = list(map(int, result))

    for i in range(len(result)):
        if result[i] == 0:
            result[i] = 2
    modified_file = BytesIO()
    file.docx.save(modified_file)
    modified_file.seek(0)
    results = CheckResults(result, modified_file)
    return results

# =========================== UTILS ================================================
def correctness_fonttable(fonts_fonttable):
    """
    :param fonts_fonttable: list font name from fontTable.xml
    :return: 0 - if Times New Roman not in list,
    1 - if only Times New Roman in list, 2 - if list contains Times New Roman and other font.
    """
    if 'Times New Roman' in fonts_fonttable:
        if len(fonts_fonttable) == 1:
            return 1
        else:
            return 2
    else:
        return 0

def iter_block_items(parent: Union[Document, _Cell]):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Expected Document or _Cell instance, but instead got {type(parent)}")

    for child in parent_elm.iterchildren():
        DOMTree = parseString(child.xml)  # TODO: тут поле для оптимизаций
        data = DOMTree.documentElement
        nodelist = data.getElementsByTagName('pic:blipFill')  # TODO: а что насчет формул? Mathtype и вордовские
        if len(nodelist) >= 1:
            for pic in nodelist:
                yield "##########################Картииинка"
        elif isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_tables_context(parent):
    para = None
    for elem in iter_block_items(parent):
        if isinstance(elem, Paragraph):
            para = elem
        elif isinstance(elem, Table):
            yield (para, elem)


def iter_pics_context(parent):
    pic = None
    for elem in iter_block_items(parent):
        if isinstance(elem, str):
            pic = elem
        elif isinstance(elem, Paragraph) and pic:
            yield (elem, pic)
            pic = None
        # else:
        #     raise Exception("Элемент не параграф и не картинка")

def heading_in_file(para):
    style = para.style
    heading = ['Heading 1', 'Heading 2', 'Heading 3']
    if style.name in heading:
        return True
    elif style.name is None:
        if style.base_style.name in heading:
            return True
    return False


def shorten(s):
    return f"\"{s}\"" if len(s) < 25 else f"\"{s[:25]}...\""

# =========================== UTILS ================================================


def get_margin_is_ok(file: Document):
    """

    :param file: docx Document()
    :return: True if margins ok else False
    """
    for section in file.sections:
        if abs(section.bottom_margin.cm - 2) > 0.001 or \
                abs(section.top_margin.cm - 2) > 0.001 or \
                abs(section.left_margin.cm - 3) > 0.001 or \
                abs(section.right_margin.cm - 1) > 0.001:
            # max_par = min(len(file.paragraphs), 3)
            max_par = 1
            highlight_mistake(Mistakes.MARGIN, paragraphs=file.paragraphs[:max_par])
            return False
    return True


def tables_is_ok(doc: Document, verifiable_paras: List[str]):
    table_is_ok = False
    table_pattern = re.compile(r"Таблица \d+(\.\d+)* – .+\.")
    flag = False
    count = 0
    for elem in iter_tables_context(doc):
        count = 1
        if elem[0] and elem[0].text != "\n" and elem[0].text != "" and elem[0].text in verifiable_paras:
            flag = True
        txt = elem[0].text if elem[0] else "НЕТ ТЕКСТА"
        valid_table = re.fullmatch(table_pattern, txt)
        if valid_table:
            if flag:
                if elem[0].paragraph_format.first_line_indent == 0:
                    table_is_ok = True
                    #print(txt, " \tХорошая таблица")
                elif elem[0].paragraph_format.first_line_indent is None:
                    if elem[0].style.paragraph_format.first_line_indent == 0:
                        table_is_ok = True
                        #print(txt, " \tХорошая таблица")
                    else:
                        table_is_ok = False
                        #print(txt, " \tПлохая таблица")
                        # TODO: весь метод может не сработать, если документ начинается сразу с таблицы
                        highlight_mistake(Mistakes.TABLES, paragraph=elem[0])
                        return table_is_ok
                else:
                    table_is_ok = False
                    #print(txt, " \tПлохая таблица")
                    highlight_mistake(Mistakes.TABLES, paragraph=elem[0])
                    return table_is_ok
        else:
            if flag:
                table_is_ok = False
                #print(txt, " \tПлохая таблица")
                highlight_mistake(Mistakes.TABLES, paragraph=elem[0])
                return table_is_ok
            else:
                table_is_ok = True
                #print(txt, " \tТаблица не учитывается")
    if count == 0:
        table_is_ok = True
    return table_is_ok


def pics_is_ok(doc: Document, verifiable_paras: List[str]):
    pic_is_ok = False
    pic_pattern = re.compile(r"Рисунок \d+(\.\d+)* – .+\.")
    flag = False
    count = 0
    for elem in iter_pics_context(doc):
        count = 1
        if elem[0] and elem[0].text != "\n" and elem[0].text != "" and elem[0].text in verifiable_paras:
            flag = True
        txt = elem[0].text if elem[0] else "НЕТ ТЕКСТА"
        valid_pic = re.fullmatch(pic_pattern, txt)
        if valid_pic:
            if flag:
                if elem[0].alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    pic_is_ok = True
                    #print(txt, " \tХорошая картинка")
                elif elem[0].alignment is None:
                    if elem[0].style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        pic_is_ok = True
                        #print(txt, " \tХорошая картинка со стилем")
                else:
                    pic_is_ok = False
                    highlight_mistake(Mistakes.PICTURES, paragraph=elem[0])
                    #print(txt, " \tПлохая картинка1")
                    return pic_is_ok
        else:
            if flag:
                pic_is_ok = False
                highlight_mistake(Mistakes.PICTURES, paragraph=elem[0])
                #print(txt, " \tПлохая картинка")
                return pic_is_ok
            else:
                pic_is_ok = True
                #print(txt, " \tКартинка не учитывается")
    if count == 0:
        pic_is_ok = True

    return pic_is_ok

def get_font_is_ok(file: ExtendedDocument):
    """
    :param file_name:
    :return: True if font is ok, else returns False
    """
    styles_with_theme = get_styles_with_theme(file.xml_styles_tree)
    run_with_theme = get_theme_run(file.xml_doc_tree)
    fonttable = get_font_from_fonttable(file.xml_font_table_tree)
    font_table = correctness_fonttable(fonttable)
    theme_is_ok = get_default_font_is_ok(file.xml_theme1_tree, file.xml_doc_text)

    ok_font = 'Times New Roman'
    font_is_ok = False  # todo: переделать логику - сделать по умолчанию True и убирать при ошибках
    if font_table == 0:
        print("Ошибка в таблице стилей")
        # весь текст с неправильным шрифтом, кажется
        return False

    elif font_table == 1:
        font_is_ok = True
    else:
        if run_with_theme:
            if theme_is_ok:
                font_is_ok = True
            else:
                # print("Ошибка в run есть тема")
                # todo: пока не можем детектировать такие раны
                font_is_ok = False
                # return False
        else:
            for para in file.docx.paragraphs[file.safe_table_of_content_index:]:
                for run in para.runs:
                    if run.font.name is not None:
                        if run.font.name == ok_font:
                            font_is_ok = True
                        else:
                            highlight_mistake(Mistakes.FONTS, paragraph=para)
                            #print(f"Ошибка неправильный шрифт в run {shorten(para.text)}")
                            font_is_ok = False
                            # return False
                    elif run.style.font.name is not None:
                        style = run.style
                        if style.style_id in styles_with_theme:
                            if theme_is_ok:
                                font_is_ok = True
                            else:
                                highlight_mistake(Mistakes.FONTS, paragraph=para)
                                #print(f"Ошибка в стиле run есть тема {shorten(para.text)}")
                                font_is_ok = False
                                # return False
                        else:
                            if style.font.name == ok_font:
                                font_is_ok = True
                            else:
                                highlight_mistake(Mistakes.FONTS, paragraph=para)
                                #print(f"Ошибка в стиле run {shorten(para.text)}")
                                font_is_ok = False
                                # return False
                    # elif run.style.font.name is None:
                    else:
                        if run.style.style_id in styles_with_theme:
                            if theme_is_ok:
                                font_is_ok = True
                            else:
                                #print(f"Ошибка в стиле run есть тема {shorten(para.text)}")
                                highlight_mistake(Mistakes.FONTS, paragraph=para)
                                font_is_ok = False
                                # return False
                        else:
                            if run.style.base_style is not None and run.style.base_style.font.name is not None:
                                font = run.style.font.name
                                style = run.style
                                while font is None and style.base_style:
                                    style = style.base_style
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            font_is_ok = True
                                            # todo: check it
                                            break
                                        else:
                                            #print(f"Ошибка в родительском стиле run есть тема {shorten(para.text)}")
                                            highlight_mistake(Mistakes.FONTS, paragraph=para)
                                            font_is_ok = False
                                            # return False
                                    else:
                                        font = style.font.name
                                if font != ok_font:
                                    #print(f"Ошибка не тот шрифт в родительском стиле run {shorten(para.text)}")
                                    highlight_mistake(Mistakes.FONTS, paragraph=para)
                                    font_is_ok = False
                                    # return False
                            else:
                                style = para.style
                                font = style.font.name
                                if font is None:
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            # todo: не используется font_is_ok - исправить
                                            font_is_ok = True
                                        else:
                                            #print(f"Ошибка в стиле параграфа есть тема {shorten(para.text)}")
                                            highlight_mistake(Mistakes.FONTS, paragraph=para)
                                            font_is_ok = False
                                            # return False
                                while font is None and style.base_style:
                                    style = style.base_style
                                    if style.style_id in styles_with_theme:
                                        if theme_is_ok:
                                            font_is_ok = True
                                        else:
                                            #print(f"Ошибка в родительском стиле параграфа есть тема {shorten(para.text)}")
                                            highlight_mistake(Mistakes.FONTS, paragraph=para)
                                            font_is_ok = False
                                            # return False
                                    else:
                                        font = style.font.name
                                if font == ok_font:
                                    font_is_ok = True
                                else:
                                    #print(f"Ошибка в стиле параграфа не тот шрифт {shorten(para.text)}")
                                    highlight_mistake(Mistakes.FONTS, paragraph=para)
                                    font_is_ok = False
                                    # return False

    return font_is_ok


def alignment_is_ok(doc: ExtendedDocument):
    style_name = ["Heading 1", "Heading 2", "Heading 3"]
    pic_pattern = re.compile(r"Рисунок \d+(\.\d+)* – .+\.")
    alignment_ok = False
    for para in doc.docx.paragraphs[doc.safe_table_of_content_index:]:
        if para.text and len(para.text) != 0 and para.text != "\n":
            if para.style.name in style_name or para.style.base_style and para.style.base_style.name in style_name or re.fullmatch(pic_pattern, para.text):
                if para.alignment:
                    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment_ok = True
                        #print(f"Выравнивание заголовка или подписи к рисунку правильное {shorten(para.text)}")
                    else:
                        alignment_ok = False
                        highlight_mistake(Mistakes.ALIGNMENT, paragraph=para,
                                          additional_info="Выравнивание заголовка или подписи к рисунку неправильное")
                        #print(f"Выравнивание заголовка или подписи к рисунку неправильное {shorten(para.text)}")
                else:
                    style = para.style
                    alignment = para.style.paragraph_format.alignment
                    while alignment is None and style.base_style:
                        style = style.base_style
                        alignment = style.paragraph_format.alignment
                    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment_ok = True
                        #print(f"Выравнивание заголовка или подписи к рисунку правильное {shorten(para.text)}")
                    else:
                        alignment_ok = False
                        highlight_mistake(Mistakes.ALIGNMENT, paragraph=para,
                                          additional_info="Выравнивание заголовка или подписи к рисунку неправильное")
                        #print(f"Выравнивание заголовка или подписи к рисунку неправильное {shorten(para.text)}")
            else:
                if para.alignment:
                    if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment_ok = True
                        #print(f"Выравнивание обычного текста правильное {shorten(para.text)}")
                    else:
                        alignment_ok = False
                        highlight_mistake(Mistakes.ALIGNMENT, paragraph=para)
                        #print(f"Выравнивание обычного текста неправильное {shorten(para.text)}")
                else:
                    style = para.style
                    alignment = para.style.paragraph_format.alignment
                    while alignment is None and style.base_style:
                        style = style.base_style
                        alignment = style.paragraph_format.alignment
                    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment_ok = True
                        #print(f"Выравнивание обычного текста правильное {shorten(para.text)}")
                    else:
                        alignment_ok = False
                        highlight_mistake(Mistakes.ALIGNMENT, paragraph=para)
                        #print(f"Выравнивание обычного текста неправильное {shorten(para.text)}")
        elif para.text == "\n":
            alignment_ok = True
        if not alignment_ok:
            return alignment_ok

    return alignment_ok


def line_spacing_is_ok(doc: ExtendedDocument):
    line_spacing_ok = False

    for para in doc.docx.paragraphs[doc.safe_table_of_content_index:]:
        if para.text and para.text != "\n":
            if para.paragraph_format.line_spacing_rule:
                if para.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                    line_spacing_ok = True
                else:
                    line_spacing_ok = False
                    highlight_mistake(Mistakes.LINE_SPACING, paragraph=para)
                    # print(f"Это параграф с неправильным межстрочным интервалом: {shorten(para.text)}")
            else:
                style = para.style
                line_spacing = style.paragraph_format.line_spacing_rule
                while line_spacing is None and style.base_style:
                    style = style.base_style
                    line_spacing = style.paragraph_format.line_spacing_rule
                if line_spacing == WD_LINE_SPACING.ONE_POINT_FIVE:
                    line_spacing_ok = True
                else:
                    line_spacing_ok = False
                    highlight_mistake(Mistakes.LINE_SPACING, paragraph=para)
                    # print(f"Это параграф неправильным стилем и межстрочным интервалом: {shorten(para.text)}")
        elif para.text == "\n":
            alignment_ok = True
        if not line_spacing_ok:
            return line_spacing_ok
    return line_spacing_ok



# print(check_file("Практика_отчет — копия.docx"))

# ======================================================================


if __name__ == "__main__":
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename()
    print(file_check_interface(file_path, mock=False).json_result)
