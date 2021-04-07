from docx.document import Document
import docx
from FormatChecker.Utils.zip import get_inner_file
from FormatChecker.Utils.xml import get_xml_repr

from io import BytesIO

# class DocumentAttributes:
#     def __init__(self):
#         self.table_of_contents_index = None
#

class ExtendedDocument:
    """
    Расширение-контейнер для хранения атрибутов и нескольких представлений документа
    """
    def __init__(self, docx_file):
        self.os_file = docx_file
        self.docx: Document = docx.Document(self.os_file)
        self._xml_theme1_path = "word/theme/theme1.xml"
        self._xml_doc_path = "word/document.xml"
        self._xml_styles_path = "word/styles.xml"
        self._xml_font_table_path = "word/fontTable.xml"

        # pure files
        self.xml_theme1_file = None
        self.xml_doc_file = None
        self.xml_styles_file = None
        self.xml_font_table_file = None

        # xml trees
        self.xml_theme1_tree = None
        self.xml_doc_tree = None
        self.xml_styles_tree = None
        self.xml_font_table_tree = None

        # plain texts
        self.xml_theme1_text = None
        self.xml_doc_text = None
        self.xml_styles_text = None
        self.xml_font_table_text = None

        self.table_of_content_index = None
        self.safe_table_of_content_index = 0

        self.__read_xmls__()
        self.table_of_content_index = self.__get_ind_content()
        if self.table_of_content_index is not None:
            self.safe_table_of_content_index = self.table_of_content_index

    def __read_xmls__(self):
        """
        Читает все нужные xml из файла docx и сохраняет в себя
        :return: None
        """
        files = get_inner_file(self.os_file, inner_paths={
            "xml_theme1": self._xml_theme1_path,
            "xml_doc": self._xml_doc_path,
            "xml_styles": self._xml_styles_path,
            "xml_font_table": self._xml_font_table_path
        })
        self.xml_theme1_file = BytesIO(files["xml_theme1"])
        self.xml_doc_file = BytesIO(files["xml_doc"])
        self.xml_styles_file = BytesIO(files["xml_styles"])
        self.xml_font_table_file = BytesIO(files["xml_font_table"])

        self.xml_theme1_tree = get_xml_repr(self.xml_theme1_file)
        self.xml_doc_tree = get_xml_repr(self.xml_doc_file)
        self.xml_styles_tree = get_xml_repr(self.xml_styles_file)
        self.xml_font_table_tree = get_xml_repr(self.xml_font_table_file)

        self.xml_theme1_text = str(self.xml_theme1_file)
        self.xml_doc_text = str(self.xml_doc_file)
        self.xml_styles_text = str(self.xml_styles_file)
        self.xml_font_table_text = str(self.xml_font_table_file)

    def __get_ind_content(self):
        """
        Получает индекс параграфа с автосодержанием в документе
        :return: None если не найден, иначе int
        """
        ind_para_content = None
        for i in range(len(self.docx.paragraphs)):
            if self.docx.paragraphs[i].text != "Содержание":
                continue
            # if i + 1 < len(self.docx.paragraphs):
            if self.docx.paragraphs[i + 1].style.name == "toc 1":
                ind_para_content = i
                break


        return ind_para_content

    def white_paras(self):
        return self.__paras(self.safe_table_of_content_index, True)

    def __paras(self, offset=0, ignore_flags=False):
        append = True
        for para in self.docx.paragraphs[offset:]:
            para_text = para.text.strip().lower()
            if ignore_flags:
                if "@start" in para_text:
                    append = False
                    continue
                elif "@stop" in para_text:
                    append = True
                    continue
            if append:
                yield para




